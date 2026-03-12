"""
parser.py — Accurate .docx → HTML converter
Handles: headings, paragraphs, bullet/numbered nested lists,
         tables, hyperlinks, code blocks (Courier font), images (base64).
"""

import base64
import os
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# ── Namespaces ────────────────────────────────────────────────────────────────
R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── Helpers ───────────────────────────────────────────────────────────────────

def _esc(text: str) -> str:
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _is_courier(run_el) -> bool:
    """Return True if the run uses Courier New (code font)."""
    rPr = run_el.find(qn("w:rPr"))
    if rPr is None:
        return False
    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        return False
    ascii_font = fonts.get(qn("w:ascii"), "")
    return "Courier" in ascii_font


def _run_inline_html(run_el, is_code_block: bool = False) -> str:
    """Convert a <w:r> element to inline HTML, with bold/italic/underline."""
    rPr = run_el.find(qn("w:rPr"))

    # Collect text tokens (handle <w:br> line breaks)
    parts = []
    for child in run_el:
        tag = child.tag.split("}")[-1]
        if tag == "t":
            parts.append(_esc(child.text or ""))
        elif tag == "br":
            parts.append("\n" if is_code_block else "<br>")

    text = "".join(parts)
    if not text:
        return ""

    if is_code_block:
        return text  # raw text inside <pre><code>

    # Apply inline styles
    if rPr is not None:
        bold      = rPr.find(qn("w:b"))
        italic    = rPr.find(qn("w:i"))
        underline = rPr.find(qn("w:u"))
        if bold      is not None: text = f"<strong>{text}</strong>"
        if italic    is not None: text = f"<em>{text}</em>"
        if underline is not None and underline.get(qn("w:val"), "") != "none":
            text = f"<u>{text}</u>"
    return text


def _para_inline_html(para, rels: dict, is_code_block: bool = False) -> str:
    """Convert all inline content (runs + hyperlinks) inside a paragraph."""
    parts = []
    for child in para._p:
        tag = child.tag.split("}")[-1]

        if tag == "r":
            parts.append(_run_inline_html(child, is_code_block))

        elif tag == "hyperlink":
            r_id = child.get(f"{{{R_NS}}}id")
            href = rels.get(r_id, "#") if r_id else "#"
            link_texts = []
            for r in child.findall(qn("w:r")):
                link_texts.append(_run_inline_html(r))
            link_text = "".join(link_texts)
            if link_text:
                parts.append(f'<a href="{_esc(href)}" target="_blank">{link_text}</a>')

    return "".join(parts)


def _table_to_html(table) -> str:
    """Convert a docx Table to a styled <table>."""
    rows_html = []
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            tag  = "th" if i == 0 else "td"
            text = _esc(" ".join(p.text for p in cell.paragraphs).strip())
            cells.append(f"<{tag}>{text}</{tag}>")
        rows_html.append("<tr>" + "".join(cells) + "</tr>")
    return (
        '<table class="doc-table">\n'
        + "<thead>" + rows_html[0] + "</thead>\n"
        + "<tbody>" + "\n".join(rows_html[1:]) + "</tbody>\n"
        + "</table>"
    )


def _image_to_html(para, doc) -> str:
    """Find the drawing in a paragraph and return an <img> tag (base64 embedded)."""
    for drawing in para._p.findall(".//" + qn("w:drawing")):
        for blip in drawing.findall(f".//{{{A_NS}}}blip"):
            r_id = blip.get(f"{{{R_NS}}}embed")
            if r_id and r_id in para.part.rels:
                rel = para.part.rels[r_id]
                image_part = rel.target_part
                
                # Get extension and mime type
                ext = image_part.content_type.split("/")[-1]
                mime = {"jpeg": "jpeg", "jpg": "jpeg", "png": "png", "gif": "gif"}.get(ext, "jpeg")
                
                # Encode binary data to base64
                b64 = base64.b64encode(image_part.blob).decode()
                return f'<img src="data:image/{mime};base64,{b64}" alt="Document image" style="max-width:100%;height:auto;border-radius:8px;margin:20px 0;display:block;">'
    return ""


def _get_list_info(para) -> tuple:
    """
    Returns (list_tag, indent_level) for list paragraphs.
    - 'List Bullet' / 'List Bullet 2' → <ul>, level 0 / 1
    - 'List Number' / 'List Number 2' → <ol>, level 0 / 1
    """
    style = para.style.name
    mapping = {
        "List Bullet":   ("ul", 0),
        "List Bullet 2": ("ul", 1),
        "List Bullet 3": ("ul", 2),
        "List Number":   ("ol", 0),
        "List Number 2": ("ol", 1),
        "List Number 3": ("ol", 2),
    }
    return mapping.get(style, (None, 0))


# ── Main Parser ───────────────────────────────────────────────────────────────

def docx_to_html(docx_path: str, images_dir: str = None) -> str:
    """
    Parse a .docx file and return a clean HTML body string.
    images_dir: folder containing extracted media files (optional).
    """
    if images_dir is None:
        images_dir = os.path.dirname(docx_path)

    doc = Document(docx_path)

    # Build rels map: rId → target_ref
    rels = {rid: rel.target_ref for rid, rel in doc.part.rels.items()}

    html   = []
    # list_stack: list of (tag, level) currently open
    stack  = []

    def close_stack_to(target_level: int):
        """Close list tags deeper than target_level."""
        while stack and stack[-1][1] >= target_level:
            tag, _ = stack.pop()
            html.append(f"</{tag}>")

    def close_all():
        while stack:
            tag, _ = stack.pop()
            html.append(f"</{tag}>")

    # Walk body children in document order (preserves table positions)
    body         = doc.element.body
    para_index   = 0
    table_index  = 0
    para_list    = list(doc.paragraphs)
    table_list   = list(doc.tables)

    for child in body:
        tag = child.tag.split("}")[-1]

        # ── TABLE ─────────────────────────────────────────────────────────────
        if tag == "tbl":
            close_all()
            if table_index < len(table_list):
                html.append(_table_to_html(table_list[table_index]))
                table_index += 1

        # ── PARAGRAPH ─────────────────────────────────────────────────────────
        elif tag == "p":
            if para_index >= len(para_list):
                para_index += 1
                continue
            para       = para_list[para_index]
            para_index += 1
            style      = para.style.name
            inline     = _para_inline_html(para, rels)
            raw_text   = para.text.strip()

            # ── Heading ───────────────────────────────────────────────────────
            if style.startswith("Heading"):
                close_all()
                try:
                    lvl = int(style.split()[-1])
                except ValueError:
                    lvl = 2
                lvl = max(1, min(6, lvl))
                html.append(f"<h{lvl}>{inline}</h{lvl}>")

            # ── List items ────────────────────────────────────────────────────
            elif style.startswith("List"):
                list_tag, indent = _get_list_info(para)
                if list_tag is None:
                    # fallback — treat as bullet
                    list_tag, indent = "ul", 0

                if not stack:
                    stack.append((list_tag, indent))
                    html.append(f"<{list_tag}>")
                else:
                    cur_tag, cur_level = stack[-1]
                    if indent > cur_level:
                        # go deeper
                        stack.append((list_tag, indent))
                        html.append(f"<{list_tag}>")
                    elif indent < cur_level:
                        # come back up
                        close_stack_to(indent)
                        if not stack or stack[-1][1] != indent:
                            stack.append((list_tag, indent))
                            html.append(f"<{list_tag}>")
                    # same level — just add item (tag change is ignored; valid HTML)

                html.append(f"<li>{inline}</li>")

            # ── Code block (Courier font) ──────────────────────────────────────
            elif _is_courier(para._p.find(qn("w:r"))) if para._p.find(qn("w:r")) is not None else False:
                close_all()
                code_text = _para_inline_html(para, rels, is_code_block=True)
                html.append(f"<pre><code>{code_text.strip()}</code></pre>")

            # ── Image paragraph ───────────────────────────────────────────────
            elif para._p.findall(".//" + qn("w:drawing")):
                close_all()
                img_html = _image_to_html(para, doc)
                if img_html:
                    html.append(img_html)

            # ── Caption ───────────────────────────────────────────────────────
            elif style == "Caption":
                close_all()
                if raw_text:
                    html.append(f'<p class="caption"><em>{inline}</em></p>')

            # ── Regular paragraph ─────────────────────────────────────────────
            elif raw_text:
                close_all()
                html.append(f"<p>{inline}</p>")

    close_all()
    return "\n".join(html)
